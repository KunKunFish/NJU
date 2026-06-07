from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/yucongyuan/Desktop/NJU")
SOURCE = ROOT / "中期报告-符合写作提纲.md"
OUTPUT = ROOT / "中期报告-按规范整理.docx"
TEMPLATE_COVER = Path("/private/tmp/template_docx_render/page-1.png")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if "\n" not in text and len(text) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def apply_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0.74)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.first_line_indent = Cm(0)


def add_cover(doc: Document) -> None:
    if not TEMPLATE_COVER.exists():
        raise FileNotFoundError(f"Template cover image not found: {TEMPLATE_COVER}")

    cover_section = doc.sections[0]
    cover_section.top_margin = Cm(0)
    cover_section.bottom_margin = Cm(0)
    cover_section.left_margin = Cm(0)
    cover_section.right_margin = Cm(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    p.add_run().add_picture(str(TEMPLATE_COVER), width=Cm(21), height=Cm(29.65))

    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    body_section.page_width = Cm(21)
    body_section.page_height = Cm(29.7)
    body_section.top_margin = Cm(2.54)
    body_section.bottom_margin = Cm(2.54)
    body_section.left_margin = Cm(2.54)
    body_section.right_margin = Cm(2.54)


def add_title(doc: Document, title: str, level: int) -> None:
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run(title)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)


def flush_table(doc: Document, table_lines: list[str]) -> None:
    if not table_lines:
        return
    rows = []
    for line in table_lines:
        if re.fullmatch(r"\|(?:\s*-+\s*\|)+", line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            set_cell_text(table.cell(r, c), cell_text, bold=(r == 0))
            if r == 0:
                set_cell_shading(table.cell(r, c), "D9EAF7")
    doc.add_paragraph()


def add_footer(doc: Document) -> None:
    for section in doc.sections[1:]:
        section.footer.is_linked_to_previous = False
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("南京大学工程管理硕士中期检查报告")
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(9)


def build() -> None:
    doc = Document()
    apply_doc_defaults(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    table_lines: list[str] = []
    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_table(doc, table_lines)
            table_lines = []
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(stripped)
            continue

        flush_table(doc, table_lines)
        table_lines = []

        if stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            add_title(doc, stripped[3:].strip(), 1)
            continue
        if stripped.startswith("### "):
            add_title(doc, stripped[4:].strip(), 2)
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.hanging_indent = Cm(0)
            body = p.add_run(stripped[2:].strip())
            body.font.name = "Times New Roman"
            body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            continue

        add_paragraph(doc, stripped)

    flush_table(doc, table_lines)
    add_footer(doc)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
