from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/yucongyuan/Desktop/NJU")
SOURCE = ROOT / "阶段性论文稿-约60%.md"
OUTPUT = ROOT / "阶段性论文稿-约60%.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, east_asia: str = "宋体", size: int = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 18 and "\n" not in text else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, "宋体", 10 if bold else 10, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def apply_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0.74)

    for name, east_asia, size in [("Heading 1", "黑体", 16), ("Heading 2", "黑体", 14), ("Heading 3", "黑体", 12)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.first_line_indent = Cm(0)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("专业学位（硕士）论文阶段性稿")
    set_run_font(run, "黑体", 20, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(18)
    run2 = p2.add_run("C公司充电桩选址问题研究")
    set_run_font(run2, "黑体", 18, bold=True)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(36)
    run3 = p3.add_run("当前完成度约60%")
    set_run_font(run3, "宋体", 12, bold=False)

    info = [
        ("学号", "532024150155"),
        ("作者姓名", "于琮源"),
        ("专业名称", "工程管理"),
        ("研究方向", "项目管理"),
        ("指导教师", "赵佳宝"),
        ("日期", "2026年6月"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(10)
    for idx, (k, v) in enumerate(info):
        set_cell_text(table.cell(idx, 0), k, bold=True)
        set_cell_text(table.cell(idx, 1), v)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("C公司充电桩选址问题研究（阶段性论文稿）")
        set_run_font(run, "宋体", 9)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)


def flush_table(doc: Document, table_lines: list[str]) -> None:
    if not table_lines:
        return
    rows = []
    for line in table_lines:
        if re.fullmatch(r"\|(?:\s*-+\s*\|)+", line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            set_cell_text(table.cell(r, c), cell_text, bold=(r == 0))
            if r == 0:
                set_cell_shading(table.cell(r, c), "D9EAF7")
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    apply_doc_defaults(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    table_lines: list[str] = []

    for raw in lines:
        stripped = raw.strip()
        if not stripped:
            flush_table(doc, table_lines)
            table_lines = []
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(stripped)
            continue

        flush_table(doc, table_lines)
        table_lines = []

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(18)
            run = p.add_run(stripped[2:].strip())
            set_run_font(run, "黑体", 18, bold=True)
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            p.add_run(stripped[3:].strip())
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            p.add_run(stripped[4:].strip())
            continue

        if stripped.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            p.add_run(stripped[5:].strip())
            continue

        add_paragraph(doc, stripped)

    flush_table(doc, table_lines)
    add_footer(doc)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
